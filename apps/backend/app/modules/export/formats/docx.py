from io import BytesIO

from docx import Document
from docx.shared import Pt

from app.modules.export.payload import ExportSessionData
from app.modules.export.time_format import format_clock_timestamp


def render_docx(data: ExportSessionData) -> bytes:
    document = Document()
    document.add_heading(data.display_title, level=0)

    meta = document.add_paragraph()
    meta.add_run(f"Started: {data.started_at.strftime('%Y-%m-%d %H:%M')}").font.size = Pt(10)
    if data.duration_ms is not None:
        document.add_paragraph(f"Duration: {format_clock_timestamp(data.duration_ms)}")
    if data.device_name:
        document.add_paragraph(f"Device: {data.device_name}")

    summary = data.summary_text
    if summary:
        document.add_heading("Summary", level=1)
        document.add_paragraph(summary)

    document.add_heading("Transcript", level=1)
    if not data.segments:
        document.add_paragraph("No transcript segments.")
    else:
        table = document.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        header_cells = table.rows[0].cells
        header_cells[0].text = "Speaker"
        header_cells[1].text = "Time"
        header_cells[2].text = "Text"

        for segment in data.segments:
            row = table.add_row().cells
            row[0].text = segment.speaker_label
            row[1].text = format_clock_timestamp(segment.start_ms)
            row[2].text = segment.text

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()
